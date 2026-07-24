import os
import re
import json
import uuid
import shutil
import logging
import asyncio
from datetime import datetime
from typing import List, Optional, Dict, Any
from concurrent.futures import ThreadPoolExecutor

import uvicorn
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

import cv2
import numpy as np
from PIL import Image
import easyocr
from docx import Document
from docx.shared import Inches
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

# ---------------------------------------------------------------------------
# Setup
# ---------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("KlikBCAProcessor")

app = FastAPI(title="KlikBCA Transfer Proof Processor")

# ---------------------------------------------------------------------------
# Global EasyOCR Singleton – initialized ONCE at startup
# ---------------------------------------------------------------------------
logger.info("Initializing EasyOCR reader (first run downloads ~30-50 MB model)…")
reader = easyocr.Reader(['id', 'en'], gpu=False)
logger.info("EasyOCR reader initialized successfully.")

# Thread pool for offloading CPU‑bound OCR work from the async event‑loop
executor = ThreadPoolExecutor(max_workers=4)

# Temporary directory for session data
TEMP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_runs")
os.makedirs(TEMP_DIR, exist_ok=True)

# Serve temp files (cropped previews, docx downloads)
app.mount("/temp", StaticFiles(directory=TEMP_DIR), name="temp")


# ===================================================================
# SERVE FRONTEND
# ===================================================================
@app.get("/", response_class=HTMLResponse)
async def get_index():
    index_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "index.html")
    with open(index_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


# ===================================================================
# IMAGE PREPROCESSING (BOTTOM TABLE AREA)
# ===================================================================

def preprocess_table_region(img: np.ndarray) -> np.ndarray:
    """
    Crop and enhance the bottom 35 % of the image (where the
    "Pelaksana Transaksi" table lives) for sharper OCR on digits
    and colon/dot separators.

    Pipeline:
      1. Crop bottom 35 %
      2. Rescale 2.5× (INTER_CUBIC)
      3. Grayscale
      4. CLAHE contrast enhancement (makes ':' and digits crisp)
    """
    h, w = img.shape[:2]
    crop_y = int(h * 0.65)                       # keep bottom 35 %
    table = img[crop_y:, :]

    # 2.5× upscale
    scaled = cv2.resize(table, None, fx=2.5, fy=2.5,
                        interpolation=cv2.INTER_CUBIC)

    # Grayscale → CLAHE
    gray = cv2.cvtColor(scaled, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)

    return enhanced


# ===================================================================
# LOW-LEVEL HELPERS
# ===================================================================

def normalize_time_text(text: str) -> str:
    """Replace common OCR misreads of ':' (dot, comma, semicolon,
    space, dash) with a colon.  Uses single-char matching so that
    '17.23.15' correctly becomes '17:23:15'."""
    # Single separator between digits (handles . , ; - and one space)
    return re.sub(r'(\d)[\s\.\,\;\-](\d)', r'\1:\2', text)


def _try_parse_time_only(text: str) -> Optional[str]:
    """Extract HH:MM:SS from *text* using flexible separators.
    Accepts  :  .  ;  ,  -  or single space between digit groups.
    Returns normalised 'HH:MM:SS' or None."""
    # Directly match with flexible separators (no pre-normalisation needed)
    m = re.search(r'(\d{1,2})[\s\.:;,\-](\d{2})[\s\.:;,\-](\d{2})', text)
    if not m:
        return None
    try:
        hh, mm, ss = int(m.group(1)), int(m.group(2)), int(m.group(3))
    except ValueError:
        return None
    if not (0 <= hh < 24 and 0 <= mm < 60 and 0 <= ss < 60):
        return None
    return f"{hh:02d}:{mm:02d}:{ss:02d}"


def _try_parse_date_only(text: str) -> Optional[str]:
    """Extract DD/MM/YYYY from *text*, or None."""
    m = re.search(r'(\d{2}/\d{2}/\d{4})', text)
    return m.group(1) if m else None


def _try_parse_combined(text: str) -> Optional[Dict[str, str]]:
    """Extract 'DD/MM/YYYY HH:MM:SS' from a single text segment.

    Accepts ANY single-char separator for the time part:
      08/07/2026 17:23:15   ← standard colon
      08/07/2026 17.23.15   ← dots  (common EasyOCR output)
      08/07/2026 19.47:05   ← mixed dot + colon
      08/07/2026 17 23 15   ← spaces
    """
    m = re.search(
        r'(\d{2})/(\d{2})/(\d{4})'
        r'\s+'
        r'(\d{1,2})[\s\.:;,\-](\d{2})[\s\.:;,\-](\d{2})',
        text,
    )
    if not m:
        return None
    day, month, year = m.group(1), m.group(2), m.group(3)
    try:
        hh, mm, ss = int(m.group(4)), int(m.group(5)), int(m.group(6))
    except ValueError:
        return None
    if not (0 <= hh < 24 and 0 <= mm < 60 and 0 <= ss < 60):
        return None
    date_s = f"{day}/{month}/{year}"
    time_s = f"{hh:02d}:{mm:02d}:{ss:02d}"
    return {"date": date_s, "time": time_s,
            "datetime_str": f"{date_s} {time_s}"}


# ===================================================================
# BBOX HELPER FUNCTIONS
# ===================================================================

def _bbox_y_center(bbox) -> float:
    """Vertical centre of a bounding box."""
    return (bbox[0][1] + bbox[2][1]) / 2.0


def _bbox_x_min(bbox) -> float:
    """Left edge of a bounding box."""
    return min(bbox[0][0], bbox[3][0])


def _bbox_x_max(bbox) -> float:
    """Right edge of a bounding box."""
    return max(bbox[1][0], bbox[2][0])


# ===================================================================
# SPATIAL COORDINATE MATCHING  (core fix)
# ===================================================================

def _extract_datetime_spatial(ocr_results: list) -> Optional[Dict[str, str]]:
    """
    Pair DATE boxes with TIME boxes using spatial/coordinate logic.

    EasyOCR often splits "08/07/2026 17:52:02" into two adjacent
    bounding boxes on the same horizontal line.  This function:

      1. Finds every box containing a DD/MM/YYYY date pattern.
      2. For each date box, searches for a NEIGHBOUR box that is:
         - On the same row  (Y-centre difference ≤ 20 px)
         - To the RIGHT     (neighbour X-start > date X-max − small gap)
      3. Checks if the neighbour contains 6 digits interpretable as
         HH:MM:SS (with any separator style).
      4. Merges date + time into one datetime string.

    Returns a list of {"date", "time", "datetime_str", "row_keyword"}
    dicts, one per successfully paired occurrence.
    """
    DATE_RE = re.compile(r'\d{2}/\d{2}/\d{4}')

    # Collect date-bearing boxes
    date_boxes = []
    for idx, (bbox, text, prob) in enumerate(ocr_results):
        if DATE_RE.search(text):
            date_boxes.append((idx, bbox, text))

    paired: List[Dict[str, Any]] = []

    for d_idx, d_bbox, d_text in date_boxes:
        date_str = _try_parse_date_only(d_text)
        if not date_str:
            continue

        d_y = _bbox_y_center(d_bbox)
        d_x_right = _bbox_x_max(d_bbox)

        # -- First check: is time already inside this same box? --
        combined = _try_parse_combined(d_text)
        if combined:
            # Determine which row-keyword this box belongs to
            row_kw = _find_row_keyword(ocr_results, d_bbox)
            combined["row_keyword"] = row_kw
            paired.append(combined)
            continue

        # -- Second check: look for a SEPARATE time box to the right --
        best_time_box = None
        best_x_dist = float('inf')

        for t_idx, (t_bbox, t_text, t_prob) in enumerate(ocr_results):
            if t_idx == d_idx:
                continue

            t_y = _bbox_y_center(t_bbox)
            t_x_left = _bbox_x_min(t_bbox)

            # Same row?  (Y-centre within 20 px)
            if abs(d_y - t_y) > 20:
                continue

            # To the right of the date box?
            x_gap = t_x_left - d_x_right
            if x_gap < -15:          # allow tiny overlap
                continue

            # Contains something that looks like a time?
            time_str = _try_parse_time_only(t_text)
            if not time_str:
                continue

            # Prefer the closest box to the right
            if x_gap < best_x_dist:
                best_x_dist = x_gap
                best_time_box = (t_bbox, time_str)

        if best_time_box:
            t_bbox_used, time_str = best_time_box
            row_kw = _find_row_keyword(ocr_results, d_bbox)
            paired.append({
                "date": date_str,
                "time": time_str,
                "datetime_str": f"{date_str} {time_str}",
                "row_keyword": row_kw,
            })
        else:
            # Date found but no time – still record it
            row_kw = _find_row_keyword(ocr_results, d_bbox)
            paired.append({
                "date": date_str,
                "time": None,
                "datetime_str": date_str,
                "row_keyword": row_kw,
            })

    return paired


def _find_row_keyword(ocr_results: list, ref_bbox) -> Optional[str]:
    """Check if 'Diotorisasi' or 'Dibuat' appears on the same row
    (Y-centre ≤ 20 px) as *ref_bbox*.  Returns the keyword string
    or None."""
    ref_y = _bbox_y_center(ref_bbox)
    for bbox, text, _prob in ocr_results:
        if abs(_bbox_y_center(bbox) - ref_y) > 20:
            continue
        tl = text.lower()
        if "diotorisasi" in tl or "otorisasi" in tl:
            return "diotorisasi"
        if "dibuat" in tl:
            return "dibuat"
    return None


def select_best_datetime(paired: list) -> Optional[Dict[str, str]]:
    """
    From a list of paired datetime dicts (each may have a
    'row_keyword'), pick the best one.

    Priority:
      1. 'Diotorisasi' row  (bottom row of table)
      2. 'Dibuat' row
      3. The one with the latest timestamp
    """
    with_time = [p for p in paired if p.get("time")]
    if not with_time:
        # Fall back to any entry that at least has a date
        with_time = paired

    if not with_time:
        return None

    # Priority 1: Diotorisasi
    for p in with_time:
        if p.get("row_keyword") == "diotorisasi":
            return p

    # Priority 2: Dibuat
    for p in with_time:
        if p.get("row_keyword") == "dibuat":
            return p

    # Priority 3: latest datetime
    with_time.sort(key=lambda d: d.get("datetime_str", ""), reverse=True)
    return with_time[0]


# ===================================================================
# SINGLE IMAGE PROCESSOR
# ===================================================================

def process_single_image(
    img_bytes: bytes,
    filename: str,
    session_dir: str,
    index: int,
) -> Dict[str, Any]:
    """Full pipeline: OCR → crop → spatial datetime extraction → save."""

    nparr = np.frombuffer(img_bytes, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
    if img is None:
        raise ValueError(f"Gagal membaca gambar: {filename}")

    h, w = img.shape[:2]

    # ── STEP 1: Full-image OCR (header / footer detection) ──────────
    full_results = reader.readtext(img, paragraph=False)
    for _bbox, text, prob in full_results:
        logger.info(f"[{filename}] OCR: '{text}' (p={prob:.2f})")

    # ── STEP 2: Crop boundaries ─────────────────────────────────────
    header_y = 0
    footer_y = h

    for bbox, text, _prob in full_results:
        tl = text.lower().strip()

        if "transfer dana" in tl or ("transfer" in tl and "dana" in tl):
            top_y = int(min(bbox[0][1], bbox[1][1]))
            if header_y == 0 or top_y < header_y:
                header_y = max(0, top_y - 20)

        if ("hak cipta" in tl or "undang-undang" in tl
                or "undung-undang" in tl or "dilindungi" in tl):
            bottom_y = int(max(bbox[2][1], bbox[3][1]))
            if footer_y == h or bottom_y > footer_y:
                footer_y = min(h, bottom_y + 20)

    if header_y >= footer_y:
        header_y, footer_y = 0, h

    # ── STEP 3: Save cropped image ──────────────────────────────────
    cropped_img = img[header_y:footer_y, 0:w]
    cropped_dir = os.path.join(session_dir, "cropped")
    os.makedirs(cropped_dir, exist_ok=True)
    cropped_filename = f"cropped_{index}_{filename}"
    cv2.imwrite(os.path.join(cropped_dir, cropped_filename), cropped_img)

    # ── STEP 4: Extract datetime (spatial matching) ─────────────────
    dt_result = None

    # Attempt 1 — spatial matching on full-image OCR results
    paired = _extract_datetime_spatial(full_results)
    if paired:
        dt_result = select_best_datetime(paired)
        if dt_result and dt_result.get("time"):
            logger.info(f"[{filename}] ✓ DateTime from full OCR: {dt_result['datetime_str']}")

    # Attempt 2 — enhanced table-region OCR  (ONLY if primary pass
    #             failed to capture both date AND time)
    if not dt_result or not dt_result.get("time"):
        logger.info(f"[{filename}] Primary pass incomplete → running enhanced table OCR…")
        table_enhanced = preprocess_table_region(img)
        table_results = reader.readtext(table_enhanced, paragraph=False)
        for _bbox, text, prob in table_results:
            logger.info(f"[{filename}] TableOCR: '{text}' (p={prob:.2f})")

        paired_table = _extract_datetime_spatial(table_results)
        if paired_table:
            best = select_best_datetime(paired_table)
            if best and best.get("time"):
                dt_result = best
                logger.info(f"[{filename}] ✓ DateTime from table OCR: {dt_result['datetime_str']}")
    else:
        logger.info(f"[{filename}] ⚡ Skipping enhanced table OCR (datetime already found)")

    # ── Build final result ──────────────────────────────────────────
    parsed_dt = datetime.max
    dt_str = "Tidak Terdeteksi"

    if dt_result and dt_result.get("time"):
        try:
            parsed_dt = datetime.strptime(dt_result["datetime_str"],
                                          "%d/%m/%Y %H:%M:%S")
            dt_str = dt_result["datetime_str"]
        except Exception:
            pass

    # Fallback: date-only
    if parsed_dt == datetime.max:
        fallback_date = (dt_result or {}).get("date")
        if not fallback_date:
            for _b, t, _p in full_results:
                fallback_date = _try_parse_date_only(t)
                if fallback_date:
                    break
        if fallback_date:
            try:
                parsed_dt = datetime.strptime(fallback_date, "%d/%m/%Y")
                dt_str = fallback_date + " (Jam tidak terdeteksi)"
            except Exception:
                pass

    return {
        "original_name": filename,
        "cropped_name": cropped_filename,
        "datetime_str": dt_str,
        "timestamp": parsed_dt.isoformat() if parsed_dt != datetime.max else None,
        "is_detected": parsed_dt != datetime.max,
    }


# ===================================================================
# ROUTES
# ===================================================================

@app.post("/upload")
async def upload_images(
    files: List[UploadFile] = File(...),
    session_id: str = Form(None),
):
    """Save raw uploads to disk. Returns session_id for /process SSE."""
    if not session_id or session_id in ("null", "undefined"):
        session_id = str(uuid.uuid4())

    uploads_dir = os.path.join(TEMP_DIR, session_id, "uploads")
    os.makedirs(uploads_dir, exist_ok=True)

    saved = []
    for file in files:
        content = await file.read()
        safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
        with open(os.path.join(uploads_dir, safe_name), "wb") as f:
            f.write(content)
        saved.append(safe_name)

    total = len(os.listdir(uploads_dir))
    return JSONResponse(content={
        "session_id": session_id,
        "uploaded": saved,
        "total_files": total,
    })


# ---------------------------------------------------------------
# SSE – real-time processing progress
# ---------------------------------------------------------------
@app.get("/process/{session_id}")
async def process_images_sse(session_id: str):
    """Server-Sent Events endpoint.  Processes images one by one and
    streams progress (percentage, filename, detected datetime) to the
    frontend in real time."""

    session_dir = os.path.join(TEMP_DIR, session_id)
    uploads_dir = os.path.join(session_dir, "uploads")
    if not os.path.exists(uploads_dir):
        raise HTTPException(status_code=404, detail="Session not found")

    # Manifest tracks already-processed files (supports "Add More")
    manifest_path = os.path.join(session_dir, "manifest.json")
    processed: Dict[str, Any] = {}
    if os.path.exists(manifest_path):
        with open(manifest_path, "r", encoding="utf-8") as f:
            processed = json.load(f)

    all_files = sorted(os.listdir(uploads_dir))
    unprocessed = [f for f in all_files if f not in processed]

    async def event_stream():
        total = len(all_files)
        already_done = len(processed)

        for idx, filename in enumerate(unprocessed):
            current = already_done + idx + 1

            # --- "processing" event ---
            yield (
                f"data: {json.dumps({'type': 'processing', 'current': current, 'total': total, 'filename': filename})}\n\n"
            )

            try:
                file_path = os.path.join(uploads_dir, filename)
                with open(file_path, "rb") as f:
                    img_bytes = f.read()

                # Offload heavy OCR work to thread-pool so the event-loop stays responsive
                loop = asyncio.get_event_loop()
                result = await loop.run_in_executor(
                    executor,
                    process_single_image,
                    img_bytes,
                    filename,
                    session_dir,
                    current - 1,
                )
                result["session_id"] = session_id
                result["preview_url"] = f"/temp/{session_id}/cropped/{result['cropped_name']}"

                processed[filename] = result
                with open(manifest_path, "w", encoding="utf-8") as f:
                    json.dump(processed, f)

                yield (
                    f"data: {json.dumps({'type': 'image_done', 'current': current, 'total': total, 'result': result})}\n\n"
                )

            except Exception as e:
                logger.error(f"Error processing {filename}: {e}")
                err_result: Dict[str, Any] = {
                    "original_name": filename,
                    "error": str(e),
                    "is_detected": False,
                }
                processed[filename] = err_result
                yield (
                    f"data: {json.dumps({'type': 'image_error', 'current': current, 'total': total, 'result': err_result})}\n\n"
                )

        # --- Final sorted payload ---
        all_results = list(processed.values())
        valid = sorted(
            [r for r in all_results if not r.get("error")],
            key=lambda r: r.get("timestamp") or "z",
        )
        errors = [r for r in all_results if r.get("error")]
        yield f"data: {json.dumps({'type': 'done', 'results': valid + errors})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ---------------------------------------------------------------
# Generate .docx
# ---------------------------------------------------------------
from docx.shared import Cm, Pt

@app.post("/generate-docx")
async def generate_docx(
    session_id: str = Form(...),
    ordered_images: List[str] = Form(...),
    images_per_page: int = Form(3),
):
    session_dir = os.path.join(TEMP_DIR, session_id)
    cropped_dir = os.path.join(session_dir, "cropped")
    manifest_path = os.path.join(session_dir, "manifest.json")
    
    if not os.path.exists(cropped_dir):
        raise HTTPException(status_code=404, detail="Session not found")

    # Load manifest to get OCR text (datetime)
    manifest = {}
    if os.path.exists(manifest_path):
        with open(manifest_path, "r", encoding="utf-8") as f:
            manifest = json.load(f)
            
    # Helper to find data by cropped name
    def get_data_by_cropped_name(c_name):
        for k, v in manifest.items():
            if v.get("cropped_name") == c_name:
                return v
        return None

    valid_items = []
    for img_name in ordered_images:
        safe_name = os.path.basename(img_name)
        img_path = os.path.join(cropped_dir, safe_name)
        if os.path.exists(img_path):
            valid_items.append((img_path, get_data_by_cropped_name(safe_name)))

    if not valid_items:
        raise HTTPException(status_code=400, detail="No valid images found")

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        
    chunks = [valid_items[i:i + images_per_page] for i in range(0, len(valid_items), images_per_page)]
    
    for chunk_idx, chunk in enumerate(chunks):
        if chunk_idx > 0:
            doc.add_page_break()
            
        for item_idx, (img_path, data) in enumerate(chunk):
            p = doc.add_paragraph()
            p.alignment = 1 # Center
            
            # Configure bounding box based on layout
            if images_per_page == 1:
                max_w, max_h = 19.0, 26.0
                p.paragraph_format.space_after = Pt(12)
            elif images_per_page == 2:
                max_w, max_h = 19.0, 13.0
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.space_before = Pt(8)
            else: # 3 per page
                max_w, max_h = 19.0, 8.5
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(4)
                
            r = p.add_run()
            try:
                # Calculate dimensions that fit inside bounding box
                with Image.open(img_path) as img:
                    w_px, h_px = img.width, img.height
                
                aspect = w_px / h_px
                # Try fitting by width
                target_w = max_w
                target_h = target_w / aspect
                
                # If height overflows, fit by height instead
                if target_h > max_h:
                    target_h = max_h
                    target_w = target_h * aspect
                    
                r.add_picture(img_path, width=Cm(target_w), height=Cm(target_h))
            except Exception as e:
                logger.error(f"Error adding picture to docx: {e}")
                continue
            
            # OCR data text removed as per user request

    docx_filename = f"KlikBCA_Compiled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(os.path.join(session_dir, docx_filename))

    return JSONResponse(content={
        "download_url": f"/temp/{session_id}/{docx_filename}"
    })


# ---------------------------------------------------------------
# Cleanup
# ---------------------------------------------------------------
@app.delete("/cleanup/{session_id}")
async def cleanup_session(session_id: str):
    session_dir = os.path.join(TEMP_DIR, session_id)
    if os.path.exists(session_dir):
        shutil.rmtree(session_dir)
        return {"status": "success"}
    return {"status": "not_found"}


# ===================================================================
# ENTRY POINT
# ===================================================================
if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
